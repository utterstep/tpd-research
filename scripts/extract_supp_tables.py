# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx"]
# ///
"""Extract the supplementary tables from the Visscher et al. (2025) mmc1.docx.

Outputs one CSV per table into data/published/. These are the *published* values
we will compare our independent re-derivation against (plus the per-study
moderator codings in S3.1/S3.2 that seed the study registry).

Run:  uv run scripts/extract_supp_tables.py
"""

import csv
import sys
from pathlib import Path

import docx  # python-docx

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "1-s2.0-S1747938X2500079X-mmc1.docx"
OUT = ROOT / "data" / "published"

# Table index (in document order) -> output filename stem.
# Verified against the document's section headings.
TABLES = {
    1: "S3.1_methods_units_settings",  # 143 studies x methods/units/settings codes
    2: "S3.2_treatment",               # 143 studies x treatment codes (incl. 4 principles)
    3: "S4.1_rho_sensitivity",         # null model vs assumed rho
    4: "S4.2_winsorized",              # meta-regression after winsorizing outliers
    5: "S4.3_ses_missing",             # meta-regression with SES "Not reported" category
    6: "S4.4_selection_model",         # weightr selection-model results
    7: "S5.1_confirmatory",            # confirmatory meta-regression (mirrors Table 4)
}


def dump(table, path: Path) -> int:
    with path.open("w", newline="") as f:
        w = csv.writer(f)
        for row in table.rows:
            w.writerow([c.text.strip().replace("\n", " ") for c in row.cells])
    return len(table.rows)


def main() -> int:
    if not DOCX.exists():
        sys.exit(f"Supplementary docx not found: {DOCX}")
    OUT.mkdir(parents=True, exist_ok=True)
    d = docx.Document(str(DOCX))
    for idx, stem in TABLES.items():
        if idx >= len(d.tables):
            print(f"!! table index {idx} missing (doc has {len(d.tables)} tables)")
            continue
        n = dump(d.tables[idx], OUT / f"{stem}.csv")
        print(f"wrote data/published/{stem}.csv ({n} rows)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
